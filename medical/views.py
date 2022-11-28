from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.contrib.auth import authenticate, login
from .forms import FollowUpForm, LoginForm, MedicalHistoryForm, PacientForm, PreprocedureCardForm, SurgicalHistoryForm, GastrointestinalProcedureForm, UrologicalProcedureForm, SurgicalProceduralDetailForm, RoboticArmLocationForm, TrocardLocationForm, BloodLossForm, RobotMalfunctionForm, InstrumentsUsedForm, AncillaryInstrumentsForm, PostProceduralForm
from django.contrib.auth.decorators import login_required
from .models import FollowUp, Pacient, PreprocedureCard, Card, MedicalHistory, SurgicalHistory, GastrointestinalProcedure, UrologicalProcedure, SurgicalProceduralDetail, RoboticArmLocation, TrocardLocation, BloodLoss, RobotMalfunction, InstrumentUsed, AncillaryInstruments, PostProcedural
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.db.models import Q, F
import os


def user_login(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            cd = form.cleaned_data
            user = authenticate(username=cd['username'], password=cd['password'])
            if user is not None:
                if user.is_active:
                    login(request, user)
                    return HttpResponse('Authenticated successfully')
            else:
                return HttpResponse('Invalid login')
    else:
        form = LoginForm()
    return render(request, 'login.html', {'form': form})


def home(request):
    amount_pacient = Pacient.objects.all()
    return render(request, "admin_panel.html", {"amount_pacient": len(amount_pacient)})

@login_required
def show_pacient(request, pacient_id):
        pacient = get_object_or_404(Pacient, pacient_id=pacient_id)
        form = PacientForm(instance=pacient)
        cards = Card.objects.filter(pacient_id=pacient)
        return render(request, 'pacient.html', {'pacient': pacient, 'form': form, "cards": cards, })


@login_required
def add_pacient(request):
    if request.method == 'POST':
        form = PacientForm(data=request.POST)
        if form.is_valid():
            new_pacient = form.save(commit=False)
            pacient = Pacient.objects.create(name=new_pacient.name, surname=new_pacient.surname, date_birth=new_pacient.date_birth, gender=new_pacient.gender)
            pacient.save()
            new_form = PacientForm(data=pacient)
            return redirect(f'/pacient/{pacient.pacient_id}', {'pacient': pacient})
        else:
            print(form.errors.as_text())
            return render(request,
                  'add_pacient.html',
                    {"form": form})
    else:
        form = PacientForm()
    return render(request,
                  'add_pacient.html',
                    {"form": form})

@login_required
def action_with_pacient(request, pacient_id):
    if request.method == 'POST':
        pacient = Pacient.objects.get(pacient_id=pacient_id)
        cards = Card.objects.filter(pacient_id=pacient_id)
        if "deletepacient" in request.POST:
            cards.delete()
            pacient.delete()
            return redirect("/pacient/all")
        else:
            form = PacientForm(data=request.POST)
            if form.is_valid():
                new_pacient = form.save(commit=False)
                pacient.name = new_pacient.name
                pacient.surname = new_pacient.surname
                pacient.date_birth = new_pacient.date_birth
                pacient.gender = new_pacient.gender
                pacient.save()
                cards = Card.objects.filter(pacient_id=pacient_id)
            else:
                return render(request, 'pacient.html', {'pacient': pacient, 'form': form, "cards": cards, })
    return redirect(f'/pacient/{pacient.pacient_id}', {'pacient': pacient, "cards": cards})

@login_required
def edit_pacient(request, pacient_id):
    pacient = Pacient.objects.get(pacient_id=pacient_id)
    form = PacientForm()
    return redirect(f'/pacient/{pacient.pacient_id}', {'pacient': pacient})

@login_required
def show_all_pacients(request):
    object_list = Pacient.objects.all() 
    paginator = Paginator(object_list, 4)
    page = request.GET.get('page')
    try:
        pacients = paginator.page(page)
    except PageNotAnInteger:
        pacients = paginator.page(1)
    except EmptyPage:
        pacients = paginator.page(paginator.num_pages)
    return render(request, "pacient_list.html",
                  {'page': page,
                   'pacients': pacients})

@login_required
def search_pacient(request):
    results = []
    if request.method == "GET":
        query = request.GET.get('search', None)
        if query == '':
             pacients = Pacient.objects.all()
        else:
            if not query.isdigit():
                pacients = Pacient.objects.filter(Q(name__icontains=query) | Q(surname__icontains=query))
            else:
                pacients = Pacient.objects.filter(Q(pacient_id__contains=int(query)))
        paginator = Paginator(pacients, 4)
        page = request.GET.get('page')
        try:
            pacients = paginator.page(page)
        except PageNotAnInteger:
            pacients = paginator.page(1)
        except EmptyPage:
            pacients = paginator.page(paginator.num_pages)
        return render(request, "pacient_list.html",
                  {'page': page,
                   'pacients': pacients})

@login_required
def show_card(request, card_id):
    card = Card.objects.get(card_id=card_id)
    form_preproc = PreprocedureCardForm(instance=PreprocedureCard.objects.get(card_id=card))
    form_mh = MedicalHistoryForm(instance=MedicalHistory.objects.get(card_id=card))
    form_sh = SurgicalHistoryForm(instance=SurgicalHistory.objects.get(card_id=card))
    form_gp = GastrointestinalProcedureForm(instance=GastrointestinalProcedure.objects.get(card_id=card))
    form_up = UrologicalProcedureForm(instance=UrologicalProcedure.objects.get(card_id=card))
    form_spd = SurgicalProceduralDetailForm(instance=SurgicalProceduralDetail.objects.get(card_id=card))
    form_ral = RoboticArmLocationForm(instance=RoboticArmLocation.objects.get(card_id=card))
    form_tl = TrocardLocationForm(instance=TrocardLocation.objects.get(card_id=card))
    form_bl = BloodLossForm(instance=BloodLoss.objects.get(card_id=card))
    form_rm = RobotMalfunctionForm(instance=RobotMalfunction.objects.get(card_id=card))
    form_iu = InstrumentsUsedForm(instance=InstrumentUsed.objects.get(card_id=card))
    form_ai = AncillaryInstrumentsForm(instance=AncillaryInstruments.objects.get(card_id=card))
    form_pp = PostProceduralForm(instance=PostProcedural.objects.get(card_id=card))
    form_fu = FollowUpForm(instance=FollowUp.objects.get(card_id=card))
    context = {"card": card, 
                "form_preproc": form_preproc, 
                "form_mh": form_mh,
                "form_sh": form_sh,
                "form_gp": form_gp,
                "form_up": form_up,
                "form_spd": form_spd,
                "form_ral": form_ral,
                "form_tl": form_tl,
                "form_bl": form_bl,
                "form_rm": form_rm,
                "form_iu": form_iu,
                "form_ai": form_ai,
                "form_pp": form_pp,
                "form_fu": form_fu}
    return render(request, "pacient_card.html", context)

@login_required
def add_new_card(request, pacient_id):
    card = Card(pacient_id=Pacient.objects.get(pacient_id=pacient_id))
    card.save()
    preprocedure = PreprocedureCard(card_id = card)
    preprocedure.save()
    medical_history = MedicalHistory(card_id = card)
    medical_history.save()
    surgical_history = SurgicalHistory(card_id = card)
    surgical_history.save()
    gastro_procedure = GastrointestinalProcedure(card_id = card)
    gastro_procedure.save()
    uro_procedure = UrologicalProcedure(card_id = card)
    uro_procedure.save()
    surg_proc_detail = SurgicalProceduralDetail(card_id = card)
    surg_proc_detail.save()
    rob_arm_location = RoboticArmLocation(card_id = card)
    rob_arm_location.save()
    trocard_location = TrocardLocation(card_id = card)
    trocard_location.save()
    blood_loss = BloodLoss(card_id = card)
    blood_loss.save()
    robot_malf = RobotMalfunction(card_id = card)
    robot_malf.save()
    inst_used = InstrumentUsed(card_id = card)
    inst_used.save()
    anc_inst = AncillaryInstruments(card_id = card)
    anc_inst.save()
    post_proc = PostProcedural(card_id = card)
    post_proc.save()
    follow_up = FollowUp(card_id = card)
    follow_up.save()
    return redirect(show_card, card.get_id())

@login_required
def edit_card(request, card_id):
    if "editcard" in request.POST:
        card = PreprocedureCard.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = PreprocedureCardForm(data=request.POST)
        print(form.errors)
        if form.is_valid():
            new_card = form.save(commit=False)
            card.creation_date = new_card.creation_date
            card.admission_date = new_card.admission_date
            card.sign_date = new_card.sign_date
            card.is_smoker = new_card.is_smoker
            card.packyears = new_card.packyears
            card.height = new_card.height
            card.weight = new_card.weight
            card.save()
            card = Card.objects.get(card_id=card_id)
            form = PreprocedureCardForm(instance=card)
            return redirect(show_card, card.get_id())
        else:
            card = Card.objects.get(card_id=card_id)
            return redirect(show_card, card.get_id())

@login_required
def edit_medical_history(request, card_id):
    if "editcard" in request.POST:
        card = MedicalHistory.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = MedicalHistoryForm(data=request.POST)
        if form.is_valid():
            new_card = form.save(commit=False)
            card = new_card
            card.history_card_id = MedicalHistory.objects.get(card_id=Card.objects.get(card_id=card_id)).history_card_id
            card.card_id = Card.objects.get(card_id=card_id)
            card.save()
        return redirect(show_card, card_id)

@login_required
def edit_surgical_history(request, card_id):
    if "editcard" in request.POST:
        card = SurgicalHistory.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = SurgicalHistoryForm(data=request.POST)
        if form.is_valid():
            new_card = form.save(commit=False)
            card.has_abdominal_surgery = new_card.has_abdominal_surgery
            card.surgion_description = new_card.surgion_description
            card.save()
        return redirect(show_card, card_id)

@login_required
def edit_gastro_procedure(request, card_id):
    if "editcard" in request.POST:
        card = GastrointestinalProcedure.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = GastrointestinalProcedureForm(data=request.POST)
        if form.is_valid():
            new_card = form.save(commit=False)
            card = new_card
            card.gastro_procedure_id = GastrointestinalProcedure.objects.get(card_id=Card.objects.get(card_id=card_id)).gastro_procedure_id
            card.card_id = Card.objects.get(card_id=card_id)
            card.save()
        return redirect(show_card, card_id)

@login_required
def edit_uro_procedure(request, card_id):
    if "editcard" in request.POST:
        card = UrologicalProcedure.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = UrologicalProcedureForm(data=request.POST)
        if form.is_valid():
            new_card = form.save(commit=False)
            card = new_card
            card.uro_procedure_id = UrologicalProcedure.objects.get(card_id=Card.objects.get(card_id=card_id)).uro_procedure_id
            card.card_id = Card.objects.get(card_id=card_id)
            card.save()
        return redirect(show_card, card_id)

@login_required
def edit_spd(request, card_id):
    if "editcard" in request.POST: 
        card = SurgicalProceduralDetail.objects.get(card_id=Card.objects.get(card_id=card_id))
        form = SurgicalProceduralDetailForm(data=request.POST)
        if form.is_valid():
            new_card = form.save(commit=False)
            card = new_card
            card.surg_proc_detail_id = SurgicalProceduralDetail.objects.get(card_id=Card.objects.get(card_id=card_id)).surg_proc_detail_id
            card.card_id = Card.objects.get(card_id=card_id)
            card.save()
        return redirect(show_card, card_id)

@login_required
def edit_ral(request, card_id):
    if "editcard" in request.POST: 
            card = RoboticArmLocation.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = RoboticArmLocationForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.robotic_arm_loc_id = RoboticArmLocation.objects.get(card_id=Card.objects.get(card_id=card_id)).robotic_arm_loc_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_tl(request, card_id):
    if "editcard" in request.POST: 
            card = TrocardLocation.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = TrocardLocationForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.trocard_location_id = TrocardLocation.objects.get(card_id=Card.objects.get(card_id=card_id)).trocard_location_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_bl(request, card_id):
    if "editcard" in request.POST: 
            card = BloodLoss.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = BloodLossForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.blood_loss_id = BloodLoss.objects.get(card_id=Card.objects.get(card_id=card_id)).blood_loss_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_rm(request, card_id):
    if "editcard" in request.POST: 
            card = RobotMalfunction.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = RobotMalfunctionForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.rob_mal_id = RobotMalfunction.objects.get(card_id=Card.objects.get(card_id=card_id)).rob_mal_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_iu(request, card_id):
    if "editcard" in request.POST: 
            card = InstrumentUsed.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = InstrumentsUsedForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.inst_use_id = InstrumentUsed.objects.get(card_id=Card.objects.get(card_id=card_id)).inst_use_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_ai(request, card_id):
    if "editcard" in request.POST: 
            card = AncillaryInstruments.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = AncillaryInstrumentsForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.ai_id = AncillaryInstruments.objects.get(card_id=Card.objects.get(card_id=card_id)).ai_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_pp(request, card_id):
    if "editcard" in request.POST: 
            card = PostProcedural.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = PostProceduralForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.postproc_id = PostProcedural.objects.get(card_id=Card.objects.get(card_id=card_id)).postproc_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)

@login_required
def edit_fu(request, card_id):
    if "editcard" in request.POST: 
            card = FollowUp.objects.get(card_id=Card.objects.get(card_id=card_id))
            form = FollowUpForm(data=request.POST)
            print(form.errors)
            if form.is_valid():
                new_card = form.save(commit=False)
                card = new_card
                card.fu_id = FollowUp.objects.get(card_id=Card.objects.get(card_id=card_id)).fu_id
                card.card_id = Card.objects.get(card_id=card_id)
                card.save()
            return redirect(show_card, card_id)



def download_file(request, card_id):
    import docx

    def boolean_output(value):
        if value:
            return "YES"
        return " "

    def localize_yesno(value):
        if value == "Да":
            return "YES"
        return "NO"
    
    def localize_gender(value):
        if value == "Мужчина":
            return "MALE"
        return "FEMALE"

    card = Card.objects.get(card_id = card_id)
    preproc = PreprocedureCard.objects.get(card_id = card)
    pacient = Pacient.objects.get(pacient_id = card.pacient_id.pacient_id)
    medical_history = MedicalHistory.objects.get(card_id=card)
    surg_history = SurgicalHistory.objects.get(card_id = card)
    gastro_procedure = GastrointestinalProcedure.objects.get(card_id = card)
    uro_procedure = UrologicalProcedure.objects.get(card_id = card)
    spd = SurgicalProceduralDetail.objects.get(card_id = card)
    ral = RoboticArmLocation.objects.get(card_id = card)
    tl = TrocardLocation.objects.get(card_id = card)
    bl = BloodLoss.objects.get(card_id = card)
    rm = RobotMalfunction.objects.get(card_id = card)

    filepath = os.getcwd() + "\\medical\\media\\report.docx"
    filepath_ready = os.getcwd() + "\\medical\\media\\report_ready.docx"
    filename = f"Report of admission #{card_id}.docx"
    doc = docx.Document(docx = filepath)
    
    paras = doc.paragraphs

    # Это самое худшее что я делал, извини
    for table in doc.tables:
        for i in table.rows:
            for j in i.cells:
                
                j.text = j.text.replace('[PACIENT]', str(pacient))
                j.text = j.text.replace('[CREATION_DATE]', preproc.creation_date.strftime("%d-%m-%Y"))

                j.text = j.text.replace('[DIABETES]', boolean_output(medical_history.relevant_disease))
                j.text = j.text.replace('[HYPERTENSION]', boolean_output(medical_history.has_hypertension))
                j.text = j.text.replace('[CARDIOVASCULAR]', boolean_output(medical_history.has_cardiovascular))
                j.text = j.text.replace('[CORD]', boolean_output(medical_history.has_cord))
                j.text = j.text.replace('[DIABETES]', boolean_output(medical_history.relevant_disease))
                j.text = j.text.replace('[RENAL_DISEASE]', boolean_output(medical_history.has_renal_disease))
                j.text = j.text.replace('[LIVER]', boolean_output(medical_history.has_liver_disease))
                j.text = j.text.replace('[APNEA]', boolean_output(medical_history.has_sleep_aphea))
                j.text = j.text.replace('[GERD]', boolean_output(medical_history.has_gerd))
                j.text = j.text.replace('[DEPRESSION]', boolean_output(medical_history.has_depression))
                j.text = j.text.replace('[OSTEO]', boolean_output(medical_history.has_osteoarthritis))
                j.text = j.text.replace('[CHRONIC]', boolean_output(medical_history.has_chronic_pain))
                j.text = j.text.replace('[STROKE]', boolean_output(medical_history.has_stroke))

                j.text = j.text.replace('[CHOLECYSTECTOMY]', boolean_output(gastro_procedure.cholecystectomy))
                j.text = j.text.replace('[BARIATRIC]', boolean_output(gastro_procedure.bariatric))
                j.text = j.text.replace('[ESOPHAGEAL]', boolean_output(gastro_procedure.esophageal))
                j.text = j.text.replace('[GASTRIC]', boolean_output(gastro_procedure.gastric))
                j.text = j.text.replace('[HEMIRIGHT]', boolean_output(gastro_procedure.hemicolectomy))
                j.text = j.text.replace('[LIVER_SURG]', boolean_output(gastro_procedure.liver_surgery))
                j.text = j.text.replace('[UNILATERAL]', boolean_output(gastro_procedure.hernia_unilateral))
                j.text = j.text.replace('[BILATERAL]', boolean_output(gastro_procedure.hernia_bilateral))
                j.text = j.text.replace('[VENTRAL]', boolean_output(gastro_procedure.ventral_hernia))
                j.text = j.text.replace('[FUNDO]', boolean_output(gastro_procedure.fundoplication))
                j.text = j.text.replace('[SIGMOID]', boolean_output(gastro_procedure.sigmoid_resection))
                j.text = j.text.replace('[RECTAL_SURG]', boolean_output(gastro_procedure.rectal))
                j.text = j.text.replace('[IMPLANT]', boolean_output(gastro_procedure.esophagus_implants))
                j.text = j.text.replace('[LINX]', boolean_output(gastro_procedure.linx_implants))
                j.text = j.text.replace('[INDICATION_FOR_PROCEDURE_G]', gastro_procedure.indication_for_procedure)
                j.text = j.text.replace('[PROCEDURAL_DETAIL_G]', gastro_procedure.procedural_details)
                j.text = j.text.replace('[SPECIFICATION_G]', gastro_procedure.specification)
                j.text = j.text.replace('[SPECIAL_PRESENT_G]', gastro_procedure.special_conditions_present)
                j.text = j.text.replace('[OTHER_G]', gastro_procedure.other)

                j.text = j.text.replace('[RADICAL_PROSTAT]', boolean_output(uro_procedure.radical_prostatectomy))
                j.text = j.text.replace('[LYMPH_DISS]', boolean_output(uro_procedure.lymph_dissection))
                j.text = j.text.replace('[ADRENALEC]', boolean_output(uro_procedure.adrenalectomy))
                j.text = j.text.replace('[SIMPLE_PROSTAT]', boolean_output(uro_procedure.simple_prostatectomy))
                j.text = j.text.replace('[PARTIAL_NEPHR]', boolean_output(uro_procedure.partial_nephrectomy))
                j.text = j.text.replace('[RADICAL_NEPHR]', boolean_output(uro_procedure.radical_nephrectomy))
                j.text = j.text.replace('[RADICAL_CYSTEC]', boolean_output(uro_procedure.radical_cystectomy))
                j.text = j.text.replace('[UTERER_RE]', boolean_output(uro_procedure.ureter_reimplant))
                j.text = j.text.replace('[UPJ_PYE]', boolean_output(uro_procedure.pyeloplasty_UPJ))
                j.text = j.text.replace('[OTHER_U]', uro_procedure.other)
                j.text = j.text.replace('[INDICATION_FOR_PROCEDURE_U]', uro_procedure.indication_for_procedure)
                j.text = j.text.replace('[PROCEDURAL_DETAIL_U]', uro_procedure.procedural_details)
                j.text = j.text.replace('[SPECIFICATION_U]', uro_procedure.specification)
                j.text = j.text.replace('[SPECIAL_PRESENT_U]', uro_procedure.special_conditions_present)

                j.text = j.text.replace('[SURG_START_TIME]', str(spd.surg_start_time))
                j.text = j.text.replace('[DOCK_START_TIME]', str(spd.docking_begins))
                j.text = j.text.replace('[DOCK_END_TIME]', str(spd.docking_ends))
                j.text = j.text.replace('[CONSOLE_START_TIME]', str(spd.console_start_time))
                j.text = j.text.replace('[CONSOLE_END_TIME]', str(spd.console_end_time))
                j.text = j.text.replace('[SURG_END_TIME]', str(spd.surg_end_time))
                j.text = j.text.replace('[INDICATION_FOR_PROCEDURE_SPD]', str(spd.indication_for_procedure))
                j.text = j.text.replace('[PROCEDURAL_DETAIL_SPD]', str(spd.procedural_details))
                j.text = j.text.replace('[SPECIFICATION_SPD]', str(spd.specification))
                j.text = j.text.replace('[SPECIAL_PRESENT_SPD]', str(spd.special_conditions_present))

                j.text = j.text.replace('[ARMS_NUMBER]', str(ral.arms_number))
                j.text = j.text.replace('[UNDERGO_SITUATIONS]', str(ral.undergo_situations))
                j.text = j.text.replace("[RAL_11]", str(ral.table11))
                j.text = j.text.replace("[RAL_12]", str(ral.table12))
                j.text = j.text.replace("[RAL_13]", str(ral.table13))
                j.text = j.text.replace("[RAL_14]", str(ral.table14))
                j.text = j.text.replace("[RAL_21]", str(ral.table21))
                j.text = j.text.replace("[RAL_22]", str(ral.table22))
                j.text = j.text.replace("[RAL_23]", str(ral.table23))
                j.text = j.text.replace("[RAL_24]", str(ral.table24))
                j.text = j.text.replace("[RAL_31]", str(ral.table31))
                j.text = j.text.replace("[RAL_32]", str(ral.table32))
                j.text = j.text.replace("[RAL_33]", str(ral.table33))
                j.text = j.text.replace("[RAL_34]", str(ral.table34))
                j.text = j.text.replace("[RAL_41]", str(ral.table41))
                j.text = j.text.replace("[RAL_42]", str(ral.table42))
                j.text = j.text.replace("[RAL_43]", str(ral.table43))
                j.text = j.text.replace("[RAL_44]", str(ral.table44))

                j.text = j.text.replace("[TL_11]", str(tl.table11))
                j.text = j.text.replace("[TL_12]", str(tl.table12))
                j.text = j.text.replace("[TL_13]", str(tl.table13))
                j.text = j.text.replace("[TL_14]", str(tl.table14))
                j.text = j.text.replace("[TL_21]", str(tl.table21))
                j.text = j.text.replace("[TL_22]", str(tl.table22))
                j.text = j.text.replace("[TL_23]", str(tl.table23))
                j.text = j.text.replace("[TL_24]", str(tl.table24))
                j.text = j.text.replace("[TL_31]", str(tl.table31))
                j.text = j.text.replace("[TL_32]", str(tl.table32))
                j.text = j.text.replace("[TL_33]", str(tl.table33))
                j.text = j.text.replace("[TL_34]", str(tl.table34))
                j.text = j.text.replace("[TL_41]", str(tl.table41))
                j.text = j.text.replace("[TL_42]", str(tl.table42))
                j.text = j.text.replace("[TL_43]", str(tl.table43))
                j.text = j.text.replace("[TL_44]", str(tl.table44))

    for para in paras:
        para.text = para.text.replace('[RELEVANT_DISEASE]', localize_yesno(medical_history.relevant_disease))
        para.text = para.text.replace('[ADMISSION_DATE]', preproc.admission_date.strftime("%d-%m-%Y %H:%M"))
        para.text = para.text.replace('[SIGN_DATE]', preproc.sign_date.strftime("%d-%m-%Y"))
        para.text = para.text.replace('[WEIGHT]', str(preproc.weight))
        para.text = para.text.replace('[HEIGHT]', str(preproc.height))
        para.text = para.text.replace('[SMOKER]', localize_yesno(preproc.is_smoker))
        if localize_yesno(preproc.is_smoker) == "NO":
            para.text = para.text.replace('[PACKYEARS]', "0")
        para.text = para.text.replace('[PACKYEARS]', str(preproc.packyears))
        para.text = para.text.replace('[BIRTH_YEAR]', pacient.date_birth.strftime("%d-%m-%Y"))
        para.text = para.text.replace('[GENDER]', localize_gender(pacient.gender))

        para.text = para.text.replace('[HAS_ABDOMINAL_DISEASE]', localize_yesno(surg_history.has_abdominal_surgery))
        para.text = para.text.replace('[SURGEON_DESCRIPTION]', surg_history.surgion_description)

        para.text = para.text.replace('[PROCEDURE_DATE]', preproc.admission_date.strftime("%d-%m-%Y %H:%M"))
        para.text = para.text.replace('[FIRST_SURGEON_G]', gastro_procedure.first_surgeon)
        para.text = para.text.replace('[SECOND_SURGEON_G]', gastro_procedure.second_surgeon)

        para.text = para.text.replace('[FIRST_SURGEON_U]',uro_procedure.first_surgeon)
        para.text = para.text.replace('[SECOND_SURGEON_U]', uro_procedure.second_surgeon)

        para.text = para.text.replace('[TABLE_HEIGHT]', str(spd.table_height))
        para.text = para.text.replace('[POSITION_PACIENT]', spd.position)
        
        para.text = para.text.replace('[LAU]', localize_yesno(tl.local_anesthesy_used))
        para.text = para.text.replace('[TYPE_DOSE]', tl.type_dose)

        para.text = para.text.replace('[BLOOD_LOSS]', str(bl.blood_loss))
        para.text = para.text.replace('[UNDERGO_LAP]', localize_yesno(bl.undergo_conversion_lap))
        para.text = para.text.replace('[UNDERGO_LAP_R]', bl.conversion_reason_lap)
        para.text = para.text.replace('[UNDERGO_OPEN]', localize_yesno(bl.undergo_conversion_open))
        para.text = para.text.replace('[UNDERGO_OPEN_R]', bl.conversion_reason_open)
        para.text = para.text.replace('[TEMP_CONV]', localize_yesno(bl.temporary_conversion))
        para.text = para.text.replace('[ROBOT_END]', localize_yesno(bl.end_by_robot))

        para.text = para.text.replace('[MALFUNCTION]', rm.malfunction)
        para.text = para.text.replace('[MALFUNCTION_COMMENT]', rm.comment)
        
        
        
        

        

    doc.save(filepath_ready)
    
    fl = open(filepath_ready, 'rb')
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    mime_type="pain/text"

    response = HttpResponse(fl, content_type=mime_type)
    response['Content-Disposition'] = f"attachment; filename={filename}"
    return response
